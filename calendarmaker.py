"""Calendar Maker, by Al Sweigart al@inventwithpython.com
Create monthly calendars, saved to a text file and fit for printing.
This code is available at https://nostarch.com/big-book-small-python-programming
Tags: short"""

import datetime
from tkinter import Tk, Canvas, IntVar, StringVar, messagebox, filedialog
from tkinter.ttk import Label, Frame, Combobox, Spinbox, Button
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# Set up the constants:
DAYS = ('Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday',
        'Friday', 'Saturday')
MONTHS = ('January', 'February', 'March', 'April', 'May', 'June', 'July',
          'August', 'September', 'October', 'November', 'December')
HOLIDAYS = {(12, 25): 'Christmas', (1, 1): 'New Year\'s', (6, 19): 'Juneteenth', (11, 11): 'Veterans Day'}

def convert_to_html(cal_list, month, year):
    cal_html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <style>
    td {{
        width: 100px;
        height: 100px;
        font-family: sans-serif;
    }}
    th {{
        width: 100px;
        height:15px;
        font-family: sans-serif;
    }}
    h1 {{
        font-family: sans-serif;
    }}
    </style>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{MONTHS[month-1]} {year} Calendar</title>
</head>
<body>
    <h1>{MONTHS[month-1]} {year}</h1>
    <table border="1">
    <tr><th>Sunday</th><th>Monday</th><th>Tuesday</th><th>Wednesday</th><th>Thursday</th><th>Friday</th><th>Saturday</th></tr>'''
    for j in range(0, 5):
        i = j*7
        cal_html += f'''<tr><td>{cal_list[i][0]}<br><br><br>{cal_list[i][1]}</td><td>{cal_list[i+1][0]}<br><br><br>{cal_list[i+1][1]}</td><td>{cal_list[i+2][0]}<br><br><br>{cal_list[i+2][1]}</td>
<td>{cal_list[i+3][0]}<br><br><br>{cal_list[i+3][1]}</td><td>{cal_list[i+4][0]}<br><br><br>{cal_list[i+4][1]}</td><td>{cal_list[i+5][0]}<br><br><br>{cal_list[i+5][1]}</td>
<td>{cal_list[i+6][0]}<br><br><br>{cal_list[i+6][1]}</td></tr>'''
    cal_html += '''</table>
</body>
</html>'''
    return cal_html

def convert_to_drawing(cal_list, month, year):
    c.delete('calshape')
    for i in range(0, 8):
        c.create_line(i*100, 0, i*100, 520, width=2, fill='black', tags='calshape')
    for i in range(0, 6):
        c.create_line(0, i*100+20, 700, i*100+20, width=2, fill='black', tags='calshape')
    c.create_line(0, 0, 700, 0, width=2, fill='black', tags='calshape')
    for i, j in enumerate(DAYS):
        c.create_text(i*100, 0, anchor='nw', text=' '+ j, font=('Helvetica', 15, 'bold'), tags='calshape')
    for i in range(5):
        for j in range(7):
            c.create_text(j*100, i*100+20, text=' ' + cal_list[i*7+j][0] + '\n\n\n ' + cal_list[i*7+j][1].replace('<br>', ''), anchor='nw', font=('Helvetica', 15), tags='calshape')
    c.move('calshape', 5, 50)
    c.create_text(5, 5, text=f'{MONTHS[month-1]} {year}', font=('Helvetica', 35, 'bold'), tags='calshape', anchor='nw')

def convert_to_docx(cal_list, month, year):
    # Create a new document
    document = Document()
    heading = document.add_heading(f'{MONTHS[month-1]} {year}', level=0)
    heading.runs[0].font.name = "Arial"
    # Add a table with 3 rows and 3 columns
    # The 'style' argument can be used to apply a built-in Word table style
    table = document.add_table(rows=6, cols=7, style='Table Grid')
    styles = document.styles
    cell_style = styles.add_style('TableCellStyle', WD_STYLE_TYPE.PARAGRAPH)
    cell_style.font.name = 'Arial'
    cell_style.font.size = Pt(10)

    # Apply the style to all paragraphs in the table cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = cell_style
    for i in range(5):
        for j in range(7):
            table.cell(i+1, j).paragraphs[0].add_run(cal_list[i*7+j][0] + '\n\n\n' + cal_list[i*7+j][1].replace('<br>', ''))
    for i, j in enumerate(DAYS):
        table.cell(0, i).paragraphs[0].add_run(j).bold = True
    for i in range(7):
        for cell in table.columns[i].cells:
            cell.width = Pt(100)
    for i in range(1, 6):
        table.rows[i].height = Pt(100)
    table.rows[0].height = Pt(15)
    table.allow_autofit = False
    return document

def getCalendarFor(year, month):
    cal_list = []

    # Get the first date in the month. (The datetime module handles all
    # the complicated calendar stuff for us here.)
    currentDate = datetime.date(year, month, 1)

    # Roll back currentDate until it is Sunday. (weekday() returns 6
    # for Sunday, not 0.)
    while currentDate.weekday() != 6:
        currentDate -= datetime.timedelta(days=1)
    inside_month = False
    while True:  # Loop over each week in the month.

        # dayNumberRow is the row with the day number labels:
        dayNumberRow = ''
        for i in range(7):
            if currentDate.day == 1:
                inside_month = not inside_month
            dayNumberLabel = str(currentDate.day)
            dayNumberRow += '|' + dayNumberLabel + (' ' * 8)
            if inside_month:
                cal_list.append((dayNumberLabel, HOLIDAYS.get((month, currentDate.day), '<br>')))
            else:
                cal_list.append((dayNumberLabel, '<br>'))
            currentDate += datetime.timedelta(days=1) # Go to next day.

        # Check if we're done with the month:
        if currentDate.month != month:
            break

    return cal_list

def export_html():
    month = MONTHS.index(month_str.get()) + 1
    try:
        year = year_var.get()
    except:
        messagebox.showwarning(message='Please enter a year, like 1967 or 2021.')
        return
    calText = getCalendarFor(year, month)
    cal_html = convert_to_html(calText, year=year, month=month)
    filename = filedialog.asksaveasfilename(title='Select HTML file name', filetypes=(('HTML files', '*.html'),))
    if filename:
        with open(filename, 'w') as f:
            f.write(cal_html)

def export_docx():
    month = MONTHS.index(month_str.get()) + 1
    try:
        year = year_var.get()
    except:
        messagebox.showwarning(message='Please enter a year, like 1967 or 2021.')
        return
    calText = getCalendarFor(year, month)
    document = convert_to_docx(calText, year=year, month=month)
    filename = filedialog.asksaveasfilename(title='Select Word file name', filetypes=(('Word files', '*.docx'),))
    if filename:
        document.save(filename)

def show_cal():
    month = MONTHS.index(month_str.get()) + 1
    try:
        year = year_var.get()
    except:
        messagebox.showwarning(message='Please enter a year, like 1967 or 2021.')
        return
    calText = getCalendarFor(year, month)
    convert_to_drawing(calText, year=year, month=month)

root = Tk()
root.title('Calendar Maker')
f = Frame(root)
f.grid(sticky='nsew')
Label(f, text='Select a month:').grid(column=0, row=0, columnspan=2)
month_str = StringVar(value='January')
Combobox(f, textvariable=month_str, values=['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December'], state='readonly').grid(column=0, row=1, columnspan=2)
Label(f, text='Select a year:').grid(column=0, row=2, columnspan=2)
year_var = IntVar(value=2025)
Spinbox(f, textvariable=year_var, from_=0, to=10000).grid(column=0, row=3, columnspan=2)
Button(f, text='Show calendar', command=show_cal).grid(column=0, row=4)
c = Canvas(f, width=(100*7)+10, height=(100*5)+25+45)
c.grid(column=0, row=6, columnspan=2)
Button(f, text='Export as HTML', command=export_html).grid(column=1, row=4)
Button(f, text='Export as Word (.docx)', command=export_docx).grid(column=0, row=5)
root.mainloop()